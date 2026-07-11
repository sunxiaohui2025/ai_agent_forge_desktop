from __future__ import annotations

from pathlib import Path
import re
import zipfile
import xml.etree.ElementTree as ET


def _ext(path: Path) -> str:
    return path.suffix.lower().lstrip(".")


def build_office_preview(path: str | Path, name: str | None = None) -> dict:
    p = Path(path)
    ext = _ext(p)
    base = {
        "name": name or p.name,
        "ext": ext,
        "local_path": str(p),
    }
    if ext == "docx":
        return {**base, "kind": "word", "document": _preview_docx(p)}
    if ext == "pptx":
        return {**base, "kind": "ppt", "slides": _preview_pptx(p)}
    if ext == "xlsx":
        return {**base, "kind": "excel", "sheets": _preview_xlsx(p)}
    return {
        **base,
        "kind": "unsupported",
        "message": "该 Office 格式暂不支持轻量预览，请使用本地办公软件打开。",
    }


def _preview_docx(path: Path) -> dict:
    from docx import Document

    doc = Document(str(path))
    blocks: list[dict] = []
    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if text:
            style = (p.style.name if p.style else "") or ""
            blocks.append({"type": "paragraph", "text": text, "style": style})
    for table in doc.tables:
        rows: list[list[str]] = []
        for row in table.rows[:80]:
            rows.append([(cell.text or "").strip() for cell in row.cells[:12]])
        if rows:
            blocks.append({"type": "table", "rows": rows})
    return {"blocks": blocks[:260], "truncated": len(blocks) > 260}


def _preview_pptx(path: Path) -> list[dict]:
    text_re = re.compile(r"\{[^}]*\}t")
    slides: list[dict] = []
    with zipfile.ZipFile(str(path)) as zf:
        slide_names = sorted(
            (n for n in zf.namelist() if re.match(r"ppt/slides/slide\d+\.xml$", n)),
            key=lambda n: int(re.search(r"slide(\d+)\.xml", n).group(1)),
        )
        for index, part_name in enumerate(slide_names, 1):
            try:
                root = ET.fromstring(zf.read(part_name))
            except ET.ParseError:
                continue
            texts = [(el.text or "").strip() for el in root.iter() if text_re.match(el.tag) and el.text]
            texts = [t for t in texts if t]
            title = texts[0] if texts else f"第 {index} 页"
            slides.append({
                "index": index,
                "title": title,
                "lines": texts[:120],
                "truncated": len(texts) > 120,
            })
    return slides


def _preview_xlsx(path: Path) -> list[dict]:
    from openpyxl import load_workbook

    wb = load_workbook(str(path), read_only=True, data_only=True)
    sheets: list[dict] = []
    for ws in wb.worksheets:
        rows: list[list[str]] = []
        for row in ws.iter_rows(max_row=120, max_col=30, values_only=True):
            cells = ["" if value is None else str(value) for value in row]
            if any(cells):
                rows.append(cells)
        sheets.append({
            "name": ws.title,
            "rows": rows,
            "truncated": bool(ws.max_row and ws.max_row > 120),
        })
    return sheets
