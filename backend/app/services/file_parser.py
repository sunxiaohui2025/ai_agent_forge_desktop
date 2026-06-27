"""File parsing pipeline.

Routing:
- Text-like (.txt/.md/.csv/.json/.log/.yaml/.xml/.html/.htm) → direct read
- Doc/binary (.pdf/.docx/.pptx/.xlsx/.png/.jpg/...) → MinerU when configured;
  fall back to local libraries on failure or when MinerU is disabled.

Result: writes parsed_markdown / parsed_chars / parse_status / parse_engine /
parse_error / parsed_at on the UploadedFile row. Stores the FULL parsed text;
per-Agent length limits are applied at injection time in agent_runner so the
same file can be sent in full to one Agent and clipped for another.
"""
from __future__ import annotations
import asyncio
import logging
from datetime import datetime, timezone
from pathlib import Path
from typing import Callable, Awaitable

from sqlalchemy import select

from ..db.models import UploadedFile
from ..db.session import SessionLocal
from .mineru_client import MinerUClient, MinerUError

logger = logging.getLogger(__name__)

TEXT_EXTS = {".txt", ".md", ".markdown", ".csv", ".json", ".log",
             ".yaml", ".yml", ".xml", ".html", ".htm", ".svg",
             ".tsv", ".sh", ".sql"}

# Extensions MinerU handles well (per their docs)
MINERU_EXTS = {".pdf", ".docx", ".doc", ".pptx", ".ppt", ".xlsx", ".xls",
               ".png", ".jpg", ".jpeg", ".bmp", ".gif", ".webp", ".tiff"}


def _truncate(text: str, hard_limit: int) -> str:
    """Head 60% + omission marker + tail 40%. hard_limit<=0 disables truncation."""
    if hard_limit <= 0 or len(text) <= hard_limit:
        return text
    head_n = int(hard_limit * 0.6)
    tail_n = hard_limit - head_n - 64
    omitted = len(text) - head_n - tail_n
    head = text[:head_n]
    tail = text[-tail_n:] if tail_n > 0 else ""
    return f"{head}\n\n... (中间 {omitted} 字符省略) ...\n\n{tail}"


# Public alias used by agent_runner at prompt-injection time.
clip_for_prompt = _truncate


# ---------- Text & local fallbacks (no external service) ----------

def _parse_text_file(path: Path) -> str:
    """Read a plain-text file with a permissive encoding."""
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        return path.read_text(encoding="utf-8", errors="replace")


def _parse_pdf_local(path: Path) -> str:
    from pypdf import PdfReader
    reader = PdfReader(str(path))
    parts: list[str] = []
    for i, page in enumerate(reader.pages, 1):
        try:
            t = page.extract_text() or ""
        except Exception:
            t = ""
        parts.append(f"## 第 {i} 页\n\n{t.strip()}")
    return "\n\n".join(parts)


def _parse_docx_local(path: Path) -> str:
    from docx import Document
    doc = Document(str(path))
    parts: list[str] = []
    for p in doc.paragraphs:
        s = (p.text or "").strip()
        if s:
            parts.append(s)
    for t in doc.tables:
        for row in t.rows:
            parts.append(" | ".join((c.text or "").strip() for c in row.cells))
    return "\n\n".join(parts)


def _parse_xlsx_local(path: Path) -> str:
    from openpyxl import load_workbook
    wb = load_workbook(str(path), read_only=True, data_only=True)
    parts: list[str] = []
    for ws in wb.worksheets:
        parts.append(f"## Sheet: {ws.title}\n")
        for row in ws.iter_rows(values_only=True):
            cells = ["" if v is None else str(v) for v in row]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _parse_pptx_local(path: Path) -> str:
    """Extract text from a .pptx file without python-pptx.

    A .pptx is a zip of XML parts; slide text lives in ``ppt/slides/slideN.xml``
    inside ``<a:t>`` elements. We pull those out in slide order. This is a
    lightweight fallback used when MinerU is unavailable — it won't recover
    layout/images, but gives the model usable text.
    """
    import re
    import zipfile
    import xml.etree.ElementTree as ET

    text_re = re.compile(r"\{[^}]*\}t")

    with zipfile.ZipFile(str(path)) as zf:
        slide_names = sorted(
            (n for n in zf.namelist() if re.match(r"ppt/slides/slide\d+\.xml$", n)),
            key=lambda n: int(re.search(r"slide(\d+)\.xml", n).group(1)),
        )
        parts: list[str] = []
        for i, name in enumerate(slide_names, 1):
            raw = zf.read(name)
            try:
                root = ET.fromstring(raw)
            except ET.ParseError:
                continue
            texts = [el.text for el in root.iter() if text_re.match(el.tag) and el.text]
            body = "\n".join(t.strip() for t in texts if t.strip())
            if body:
                parts.append(f"## 第 {i} 页\n\n{body}")
        return "\n\n".join(parts)


def _parse_html_local(path: Path) -> str:
    from bs4 import BeautifulSoup
    raw = _parse_text_file(path)
    soup = BeautifulSoup(raw, "html.parser")
    for tag in soup(["script", "style"]):
        tag.decompose()
    return soup.get_text(separator="\n").strip()


def _local_for_ext(ext: str) -> Callable[[Path], str] | None:
    e = ext.lower()
    if e in TEXT_EXTS:
        if e in (".html", ".htm"):
            return _parse_html_local
        return _parse_text_file
    if e == ".pdf":
        return _parse_pdf_local
    if e in (".docx",):
        return _parse_docx_local
    if e in (".xlsx",):
        return _parse_xlsx_local
    if e in (".pptx",):
        return _parse_pptx_local
    return None


# ---------- Public entrypoint ----------

async def parse_uploaded_file(file_id: int) -> None:
    """Parse the file in-place. Updates DB row. Safe to call as a background task."""
    async with SessionLocal() as db:
        row = (await db.execute(select(UploadedFile).where(UploadedFile.id == file_id))).scalar_one_or_none()
        if not row:
            return
        if row.parse_status == "done":
            return
        row.parse_status = "parsing"
        row.parse_error = None
        await db.commit()

        # Load MinerU config from the DB-backed system settings (desktop mode).
        # Users configure it from the Settings UI; no .env required.
        from ..api.admin.settings import get_mineru_settings_raw
        mineru = await get_mineru_settings_raw(db)

    path = Path(row.path)
    ext = path.suffix.lower()
    parsed: str | None = None
    engine: str | None = None
    err: str | None = None

    # Text → direct
    if ext in TEXT_EXTS:
        try:
            fn = _local_for_ext(ext)
            parsed = await asyncio.to_thread(fn, path) if fn else _parse_text_file(path)
            engine = "text"
        except Exception as e:  # noqa: BLE001
            err = f"local text read failed: {e}"
    else:
        # Try MinerU first when configured. Config comes from the Settings UI
        # (system_settings table); mode="disabled" or empty api_key skips it.
        if mineru.mode != "disabled" and mineru.api_key and ext in MINERU_EXTS:
            try:
                client = MinerUClient(
                    base_url=mineru.base_url,
                    api_key=mineru.api_key,
                    timeout=mineru.timeout_sec,
                )
                parsed = await client.parse(path, row.name)
                engine = f"mineru-{mineru.mode}"
            except (MinerUError, asyncio.TimeoutError, Exception) as e:  # noqa: BLE001
                logger.warning("MinerU parse failed for %s: %s", row.name, e)
                err = f"MinerU: {e}"

        # Local lib fallback
        if parsed is None:
            fn = _local_for_ext(ext)
            if fn is not None:
                try:
                    parsed = await asyncio.to_thread(fn, path)
                    engine = "local-lib"
                    err = None  # successful fallback overrides MinerU error
                except Exception as e:  # noqa: BLE001
                    err = f"{(err + ' / ') if err else ''}local-lib: {e}"
            elif err is None:
                err = f"unsupported extension: {ext}"

    # Persist result. Full parsed text is stored — per-turn truncation (if any)
    # happens at injection time in agent_runner, honoring the Agent's own
    # `parsed_content_limit` override when set.
    async with SessionLocal() as db:
        row = (await db.execute(select(UploadedFile).where(UploadedFile.id == file_id))).scalar_one()
        if parsed is not None:
            row.parsed_markdown = parsed
            row.parsed_chars = len(parsed)
            row.parse_status = "done"
            row.parse_engine = engine
            row.parse_error = None
            row.parsed_at = datetime.now(timezone.utc)
        else:
            row.parse_status = "failed"
            row.parse_error = (err or "未知错误")[:1000]
        await db.commit()
